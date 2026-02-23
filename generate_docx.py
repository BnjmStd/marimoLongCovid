"""
generate_docx.py
================
Exporta la Tabla 1 y la Tabla 2 a un único archivo Word (.docx).

Requiere:
    python-docx  (pip install python-docx)
    Los CSV ya generados por generate_table1.py y generate_table2.py.

Uso:
    python generate_table1.py   # genera pdf_reports/table1_sociodemographic.csv
    python generate_table2.py   # genera pdf_reports/table2_longcovid.csv
    python generate_docx.py     # genera pdf_reports/tables_longcovid.docx
"""

from __future__ import annotations
from pathlib import Path
from typing import Optional
import pandas as pd
from docx import Document as DocxDocument
from docx.document import Document as DocxDoc
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─────────────────────────────────────────────
# CONFIGURACIÓN
# ─────────────────────────────────────────────
CSV_TABLE1 = Path("pdf_reports/table1_sociodemographic.csv")
CSV_TABLE2 = Path("pdf_reports/table2_longcovid.csv")
OUT_DOCX   = Path("pdf_reports/tables_longcovid.docx")

SECTION_MARKER = "───"   # identifica filas de sección (encabezados de bloque)
INDENT_MARKER  = "  "    # identifica filas de sub-categoría

# Fuente y tamaños
FONT_NAME   = "Arial"
FONT_SIZE   = Pt(9)
FONT_HEADER = Pt(9)

# Colores
COLOR_SECTION_BG = "D9D9D9"   # gris claro para encabezado de sección
COLOR_HEADER_BG  = "404040"   # gris oscuro para encabezado de columnas
COLOR_HEADER_FG  = "FFFFFF"   # blanco para texto de encabezado


# ─────────────────────────────────────────────
# HELPERS DE ESTILO
# ─────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str) -> None:
    """Pone color de fondo a una celda."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_border(cell, sides=("top", "bottom", "left", "right"),
                    sz: int = 4, color: str = "808080") -> None:
    """Aplica borde a los lados indicados de una celda."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in sides:
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), str(sz))
        border.set(qn("w:color"), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)


def fmt_cell(cell, text: str, bold: bool = False, italic: bool = False,
             font_size: Optional[Pt] = None, align: str = "left",
             fg_color: Optional[str] = None) -> None:
    """Escribe texto en una celda con formato."""
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.clear()
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = FONT_NAME
    run.font.size = font_size or FONT_SIZE
    if fg_color:
        run.font.color.rgb = RGBColor.from_string(fg_color)


def set_col_width(table, col_idx: int, width_cm: float) -> None:
    """Ajusta el ancho de una columna."""
    for row in table.rows:
        row.cells[col_idx].width = Cm(width_cm)


# ─────────────────────────────────────────────
# FUNCIÓN PRINCIPAL: escribir una tabla en el doc
# ─────────────────────────────────────────────

def write_table(doc: DocxDoc, df: pd.DataFrame, title: str,
                col_widths: list[float]) -> None:
    """
    Escribe el título y la tabla en el documento Word.

    df debe tener columnas: Variable, Average / Levels, + columnas de grupo.
    """
    # Título de la tabla
    h = doc.add_heading(title, level=1)
    h.runs[0].font.name = FONT_NAME
    h.runs[0].font.size = Pt(11)
    h.runs[0].font.color.rgb = RGBColor(0, 0, 0)

    columns = list(df.columns)
    n_cols  = len(columns)

    # Crear tabla Word
    tbl = doc.add_table(rows=1, cols=n_cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.style = "Table Grid"

    # ── Fila de encabezado ────────────────────────────────────────────────
    hdr_cells = tbl.rows[0].cells
    col_labels = {
        "Variable":         "Variable",
        "Average / Levels": "Average / Levels",
        "Men":              "Men",
        "Women":            "Women",
        "Controls":         "Controls",
        "Cases":            "Cases",
        "p-value":          "p-value²",
    }
    for i, col in enumerate(columns):
        fmt_cell(hdr_cells[i], col_labels.get(col, col),
                 bold=True, align="center",
                 fg_color=COLOR_HEADER_FG, font_size=FONT_HEADER)
        set_cell_bg(hdr_cells[i], COLOR_HEADER_BG)

    # ── Filas de datos ────────────────────────────────────────────────────
    for _, row in df.iterrows():
        var   = str(row["Variable"])
        level = str(row["Average / Levels"])

        is_section   = SECTION_MARKER in level
        is_subrow    = level.startswith(INDENT_MARKER) and not is_section
        is_main_var  = var.strip() != "" and not is_section

        cells = tbl.add_row().cells

        # Columna Variable
        fmt_cell(cells[0], var, bold=is_main_var)

        # Columna Average / Levels
        display_level = level.strip() if is_section else level
        fmt_cell(cells[1], display_level,
                 bold=is_section or is_main_var,
                 italic=False)

        # Columnas de grupos (a partir de la tercera)
        for i in range(2, n_cols):
            val = str(row[columns[i]])
            fmt_cell(cells[i], val, align="center")

        # Fondos: filas de sección en gris claro
        if is_section:
            for c in cells:
                set_cell_bg(c, COLOR_SECTION_BG)

    # ── Anchos de columna ─────────────────────────────────────────────────
    for i, w in enumerate(col_widths):
        set_col_width(tbl, i, w)

    doc.add_paragraph()   # espacio entre tablas


# ─────────────────────────────────────────────
# MAIN
# ─────────────────────────────────────────────

def main() -> None:
    # Verificar que los CSV existan
    for path in (CSV_TABLE1, CSV_TABLE2):
        if not path.exists():
            raise FileNotFoundError(
                f"No se encontró {path}. "
                "Ejecuta primero generate_table1.py y generate_table2.py."
            )

    df1 = pd.read_csv(CSV_TABLE1, dtype=str, keep_default_na=False)
    df2 = pd.read_csv(CSV_TABLE2, dtype=str, keep_default_na=False)

    doc = DocxDocument()

    # Márgenes del documento (2 cm por lado)
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2)
        section.right_margin  = Cm(2)

    # ── Tabla 1: Variable(3) | Avg/Levels(4) | Men(2.5) | Women(2.5) | Controls(2.5) | Cases(2.5)
    write_table(
        doc, df1,
        title="Table 1. Sociodemographic and Clinical Characteristics by Sex and Long COVID Status",
        col_widths=[3.5, 4.5, 2.3, 2.3, 2.3, 2.3],
    )

    # Salto de página entre tablas
    doc.add_page_break()

    # ── Tabla 2: Variable(3.5) | Avg/Levels(4.5) | Controls(2.5) | Cases(2.5) | p-value(2)
    write_table(
        doc, df2,
        title="Table 2. Sociodemographic and Clinical Characteristics by Long COVID Status",
        col_widths=[3.5, 4.5, 2.5, 2.5, 2.0],
    )

    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_DOCX))
    print(f"Documento guardado en: {OUT_DOCX}")


if __name__ == "__main__":
    main()
